"""
Form Automation Tool
A GUI application for filling out post shut reports using a template document.
"""

# TODO PLACE HOLDERS IN WORD DOC TO MEET INPUT
# TODO FIND OUT HOWE TO SELF GENERATE FIELDS AS MORE ARE ADDED
# TODO MAYBE CREATE MANY FIELDS AND HAVE PYTHON DELETE UNFILLED ONES

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import platform

class PhotoEntry(ttk.Frame):
    """A frame for photo entry with description"""
    def __init__(self, parent, description_required=True):
        super().__init__(parent)
        self.description_required = description_required
        self.photo_path = None
        self.photo_obj = None
        
        # Photo preview
        self.preview_label = ttk.Label(self)
        self.preview_label.pack(side="left", padx=5)
        
        # Description
        desc_frame = ttk.Frame(self)
        desc_frame.pack(side="left", fill="both", expand=True)
        
        ttk.Label(desc_frame, text="Description:").pack(anchor="w")
        self.description = scrolledtext.ScrolledText(desc_frame, height=3, width=40)
        self.description.pack(fill="both", expand=True)
        
        # Buttons
        btn_frame = ttk.Frame(self)
        btn_frame.pack(side="right", padx=5)
        
        ttk.Button(btn_frame, text="Select Photo", command=self.select_photo).pack(pady=2)
        ttk.Button(btn_frame, text="Remove", command=self.destroy).pack(pady=2)
    
    def select_photo(self):
        """Handle photo selection"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp")]
        )
        if file_path:
            self.photo_path = file_path
            image = Image.open(file_path)
            image.thumbnail((100, 100))
            self.photo_obj = ImageTk.PhotoImage(image)
            self.preview_label.configure(image=self.photo_obj)
    
    def get_data(self):
        """Get photo entry data"""
        if self.description_required and not self.description.get("1.0", "end-1c").strip():
            raise ValueError("Photo description is required")
        
        return {
            'path': self.photo_path,
            'description': self.description.get("1.0", "end-1c")
        }

class JobSection(ttk.LabelFrame):
    """A section for job details"""
    def __init__(self, parent, title="Job Entry"):
        super().__init__(parent, text=title, padding="10")
        
        # Work Order
        ttk.Label(self, text="Work Order #:").grid(row=0, column=0, sticky="w")
        self.wo_entry = ttk.Entry(self, width=30)
        self.wo_entry.grid(row=0, column=1, sticky="w", padx=5)
        
        # Scope
        ttk.Label(self, text="Scope:").grid(row=1, column=0, sticky="w")
        self.scope_entry = ttk.Entry(self, width=50)
        self.scope_entry.grid(row=1, column=1, sticky="ew", padx=5)
        
        # Status
        status_frame = ttk.Frame(self)
        status_frame.grid(row=2, column=0, columnspan=2, sticky="w", pady=5)
        self.status_var = tk.StringVar(value="complete")
        ttk.Radiobutton(status_frame, text="Complete", variable=self.status_var, 
                       value="complete").pack(side="left", padx=5)
        ttk.Radiobutton(status_frame, text="Incomplete", variable=self.status_var,
                       value="incomplete").pack(side="left", padx=5)
        ttk.Radiobutton(status_frame, text="Further works required", variable=self.status_var,
                       value="further_works").pack(side="left", padx=5)
        
        # Date
        ttk.Label(self, text="Date Completed:").grid(row=3, column=0, sticky="w")
        self.date_entry = ttk.Entry(self, width=20)
        self.date_entry.grid(row=3, column=1, sticky="w", padx=5)
        self.date_entry.insert(0, datetime.now().strftime("%d/%m/%Y"))
        
        # Summary
        ttk.Label(self, text="Summary (What was completed):").grid(row=4, column=0, columnspan=2, sticky="w", pady=(10,0))
        self.summary_text = scrolledtext.ScrolledText(self, height=4, width=60)
        self.summary_text.grid(row=5, column=0, columnspan=2, sticky="ew", padx=5)
        
        # Problems/Delays
        ttk.Label(self, text="Problems or Delays:").grid(row=6, column=0, columnspan=2, sticky="w", pady=(10,0))
        self.problems_text = scrolledtext.ScrolledText(self, height=4, width=60)
        self.problems_text.grid(row=7, column=0, columnspan=2, sticky="ew", padx=5)
        
        # Recommendations
        ttk.Label(self, text="Recommendations:").grid(row=8, column=0, columnspan=2, sticky="w", pady=(10,0))
        self.recommendations_text = scrolledtext.ScrolledText(self, height=4, width=60)
        self.recommendations_text.grid(row=9, column=0, columnspan=2, sticky="ew", padx=5)
        
        # Photos
        photos_frame = ttk.LabelFrame(self, text="Photos", padding="5")
        photos_frame.grid(row=10, column=0, columnspan=2, sticky="ew", pady=10)
        
        self.photos_container = ttk.Frame(photos_frame)
        self.photos_container.pack(fill="x", expand=True)
        
        ttk.Button(photos_frame, text="Add Photo", 
                  command=self.add_photo).pack(pady=5)
        
        self.photos = []
        
        # Remove button
        ttk.Button(self, text="Remove Entry", 
                  command=self.destroy).grid(row=11, column=0, columnspan=2, pady=10)
    
    def add_photo(self):
        """Add a photo entry"""
        photo_entry = PhotoEntry(self.photos_container)
        photo_entry.pack(fill="x", pady=5)
        self.photos.append(photo_entry)
    
    def get_data(self):
        """Get all data from the section"""
        return {
            'work_order': self.wo_entry.get(),
            'scope': self.scope_entry.get(),
            'status': self.status_var.get(),
            'date': self.date_entry.get(),
            'summary': self.summary_text.get("1.0", "end-1c"),
            'problems': self.problems_text.get("1.0", "end-1c"),
            'recommendations': self.recommendations_text.get("1.0", "end-1c"),
            'photos': [photo.get_data() for photo in self.photos if photo.winfo_exists()]
        }

class FormFillerApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Post Shut Report Form Filler")
        self.root.geometry("1200x800")
        
        # Get available templates
        self.templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        self.template_doc = None
        
        self.setup_ui()
    
    def setup_ui(self):
        """Set up the user interface components"""
        # Create a canvas with scrollbar
        self.canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)
        
        # Configure canvas
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)
        
        # Configure mouse wheel scrolling based on platform
        if platform.system() == "Darwin":  # macOS
            self.canvas.bind_all("<MouseWheel>", self._on_mousewheel_macos)
        else:  # Windows/Linux
            self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
            # Linux also uses Button-4 and Button-5 for mouse wheel
            self.canvas.bind_all("<Button-4>", self._on_mousewheel_linux)
            self.canvas.bind_all("<Button-5>", self._on_mousewheel_linux)
        
        # Pack scrollbar and canvas
        scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        # Main container
        main_container = ttk.Frame(self.scrollable_frame, padding="10")
        main_container.pack(fill="both", expand=True)
        
        # Document Info Frame
        doc_frame = ttk.LabelFrame(main_container, text="Document Information", padding="10")
        doc_frame.pack(fill="x", pady=5)
        
        # Template selection
        template_frame = ttk.Frame(doc_frame)
        template_frame.pack(fill="x", pady=5)
        ttk.Label(template_frame, text="Template:").pack(side="left")
        self.template_var = tk.StringVar()
        self.template_combo = ttk.Combobox(template_frame, textvariable=self.template_var)
        self.template_combo['values'] = [f for f in os.listdir(self.templates_dir) if f.endswith('.docx')]
        self.template_combo.bind('<<ComboboxSelected>>', self.on_template_selected)
        self.template_combo.pack(side="left", padx=5)
        
        # Site Title
        title_frame = ttk.Frame(doc_frame)
        title_frame.pack(fill="x", pady=5)
        ttk.Label(title_frame, text="Site Title:").pack(side="left")
        self.site_title_var = tk.StringVar()
        ttk.Entry(title_frame, textvariable=self.site_title_var, width=50).pack(side="left", padx=5)
        
        # Document ID
        id_frame = ttk.Frame(doc_frame)
        id_frame.pack(fill="x", pady=5)
        ttk.Label(id_frame, text="Document ID:").pack(side="left")
        self.id_num_var = tk.StringVar()
        ttk.Entry(id_frame, textvariable=self.id_num_var, width=20).pack(side="left", padx=5)
        
        # Introduction
        ttk.Label(doc_frame, text="Introduction:").pack(anchor="w")
        self.intro_text = scrolledtext.ScrolledText(doc_frame, height=4, width=60)
        self.intro_text.pack(fill="x")
        
        # Safety Summary
        ttk.Label(doc_frame, text="Safety Summary:").pack(anchor="w")
        self.safety_text = scrolledtext.ScrolledText(doc_frame, height=4, width=60)
        self.safety_text.pack(fill="x")
        
        # Personnel Frame
        personnel_frame = ttk.LabelFrame(main_container, text="Personnel", padding="10")
        personnel_frame.pack(fill="x", pady=5)
        
        # Personnel entries dictionary to store references
        self.personnel_entries = {}
        personnel_types = [
            "Mechanical Supervisors",
            "Service Technicians",
            "Mechanical Fitters",
            "Trade Assistants",
            "Hydraulic Technicians",
            "Hydraulic Supervisors"
        ]
        
        for i, p_type in enumerate(personnel_types):
            row_frame = ttk.Frame(personnel_frame)
            row_frame.pack(fill="x", pady=2)
            
            ttk.Label(row_frame, text=f"{p_type}:").pack(side="left")
            
            count_var = tk.StringVar(value="0")
            ttk.Entry(row_frame, textvariable=count_var, width=5).pack(side="left", padx=5)
            
            ttk.Label(row_frame, text="Location:").pack(side="left", padx=5)
            location_var = tk.StringVar()
            ttk.Entry(row_frame, textvariable=location_var, width=30).pack(side="left")
            
            self.personnel_entries[p_type] = {
                "per_shift": count_var,
                "location": location_var
            }
        
        # Job Sections
        # Completed Jobs
        completed_frame = ttk.LabelFrame(main_container, text="Completed Jobs", padding="10")
        completed_frame.pack(fill="x", pady=5)
        
        self.completed_container = ttk.Frame(completed_frame)
        self.completed_container.pack(fill="x")
        
        ttk.Button(completed_frame, text="Add Completed Job", 
                  command=lambda: self.add_job_section("completed")).pack(pady=5)
        
        # Uncompleted Jobs
        uncompleted_frame = ttk.LabelFrame(main_container, text="Uncompleted Jobs", padding="10")
        uncompleted_frame.pack(fill="x", pady=5)
        
        self.uncompleted_container = ttk.Frame(uncompleted_frame)
        self.uncompleted_container.pack(fill="x")
        
        ttk.Button(uncompleted_frame, text="Add Uncompleted Job", 
                  command=lambda: self.add_job_section("uncompleted")).pack(pady=5)
        
        # Generate button
        ttk.Button(main_container, text="Generate Document", 
                  command=self.generate_document).pack(pady=10)
        
        # Initialize job section lists
        self.completed_sections = []
        self.uncompleted_sections = []
    
    def on_template_selected(self, event):
        """Handle template selection"""
        template_name = self.template_var.get()
        if template_name:
            template_path = os.path.join(self.templates_dir, template_name)
            self.template_doc = Document(template_path)
    
    def _on_mousewheel_macos(self, event):
        """Handle mouse wheel scrolling for macOS"""
        self.canvas.yview_scroll(-1 * event.delta, "units")
    
    def _on_mousewheel(self, event):
        """Handle mouse wheel scrolling for Windows"""
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    def _on_mousewheel_linux(self, event):
        """Handle mouse wheel scrolling for Linux"""
        if event.num == 4:  # scroll up
            self.canvas.yview_scroll(-1, "units")
        elif event.num == 5:  # scroll down
            self.canvas.yview_scroll(1, "units")
    
    def add_job_section(self, section_type):
        """Add a new job section"""
        if section_type == "completed":
            section = JobSection(self.completed_container)
            self.completed_sections.append(section)
            section.pack(fill="x", pady=5)
        else:
            section = JobSection(self.uncompleted_container)
            self.uncompleted_sections.append(section)
            section.pack(fill="x", pady=5)
    
    def run(self):
        """Start the application"""
        self.root.mainloop()

    def generate_document(self):
        """Generate the Word document using the template"""
        try:
            if not self.template_doc:
                messagebox.showerror("Error", "No template selected")
                return
            
            # Get save location first
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")]
            )
            if not file_path:
                return
            
            # Copy template to new location
            template_path = os.path.join(self.templates_dir, self.template_var.get())
            import shutil
            shutil.copy2(template_path, file_path)
            
            # Open the new document
            doc = Document(file_path)
            
            # Build replacements dictionary
            replacements = {
                # Document info
                "<<SITE_TITLE>>": self.site_title_var.get(),
                "<<DOCUMENT_ID>>": self.id_num_var.get(),
                "<<INTRODUCTION_TEXT>>": self.intro_text.get("1.0", "end-1c"),
                "<<SAFETY_SUMMARY>>": self.safety_text.get("1.0", "end-1c"),
                
                # Personnel table
                "<<SUPERVISOR_COUNT>>": self.personnel_entries["Mechanical Supervisors"]["per_shift"].get(),
                "<<TECHNICIAN_COUNT>>": self.personnel_entries["Service Technicians"]["per_shift"].get(),
                "<<FITTER_COUNT>>": self.personnel_entries["Mechanical Fitters"]["per_shift"].get(),
                "<<ASSISTANT_COUNT>>": self.personnel_entries["Trade Assistants"]["per_shift"].get(),
                "<<HYDRAULIC_TECH_COUNT>>": self.personnel_entries["Hydraulic Technicians"]["per_shift"].get(),
                "<<HYDRAULIC_SUPERVISOR_COUNT>>": self.personnel_entries["Hydraulic Supervisors"]["per_shift"].get(),
                "<<PERSONNEL_LOCATION>>": self.personnel_entries["Mechanical Supervisors"]["location"].get(),  # Same location for all
            }
            
            # Add uncompleted work order placeholders
            for i, section in enumerate(self.uncompleted_sections, 1):
                data = section.get_data()
                replacements.update({
                    f"<<UNCOMPLETED_WORKORDER{i}_TITLE>>": data['scope'],
                    f"<<UNCOMPLETED_WORKORDER{i}_DESCRIPTION>>": f"WO {data['work_order']} – {data['scope']}",
                    f"<<UNCOMPLETED_WORKORDER{i}_DETAIL1>>": data['summary'].split('\n')[0] if data['summary'] else "",
                    f"<<UNCOMPLETED_WORKORDER{i}_DETAIL2>>": data['summary'].split('\n')[1] if data['summary'] and len(data['summary'].split('\n')) > 1 else "",
                    f"<<UNCOMPLETED_WORKORDER{i}_DETAIL3>>": data['summary'].split('\n')[2] if data['summary'] and len(data['summary'].split('\n')) > 2 else "",
                    f"<<UNCOMPLETED_WORKORDER{i}_DETAIL4>>": data['summary'].split('\n')[3] if data['summary'] and len(data['summary'].split('\n')) > 3 else "",
                    f"<<UNCOMPLETED_WORKORDER{i}_DETAIL5>>": data['summary'].split('\n')[4] if data['summary'] and len(data['summary'].split('\n')) > 4 else "",
                })
            
            # Add completed work order placeholders
            for i, section in enumerate(self.completed_sections, 1):
                data = section.get_data()
                replacements.update({
                    f"<<COMPLETED_WORKORDER{i}_TITLE>>": data['scope'],
                    f"<<COMPLETED_WORKORDER{i}_NUMBER>>": data['work_order'],
                    f"<<COMPLETED_WORKORDER{i}_SCOPE>>": data['scope'],
                    f"<<COMPLETED_WORKORDER{i}_DATE>>": data['date'],
                    f"<<COMPLETED_WORKORDER{i}_SUMMARY1>>": data['summary'].split('\n')[0] if data['summary'] else "",
                    f"<<COMPLETED_WORKORDER{i}_SUMMARY2>>": data['summary'].split('\n')[1] if data['summary'] and len(data['summary'].split('\n')) > 1 else "",
                    f"<<COMPLETED_WORKORDER{i}_SUMMARY3>>": data['summary'].split('\n')[2] if data['summary'] and len(data['summary'].split('\n')) > 2 else "",
                    f"<<COMPLETED_WORKORDER{i}_SUMMARY4>>": data['summary'].split('\n')[3] if data['summary'] and len(data['summary'].split('\n')) > 3 else "",
                    f"<<COMPLETED_WORKORDER{i}_PROBLEM1>>": data['problems'].split('\n')[0] if data['problems'] else "",
                    f"<<COMPLETED_WORKORDER{i}_PROBLEM2>>": data['problems'].split('\n')[1] if data['problems'] and len(data['problems'].split('\n')) > 1 else "",
                    f"<<COMPLETED_WORKORDER{i}_PROBLEM3>>": data['problems'].split('\n')[2] if data['problems'] and len(data['problems'].split('\n')) > 2 else "",
                    f"<<COMPLETED_WORKORDER{i}_RECOMMENDATION1>>": data['recommendations'].split('\n')[0] if data['recommendations'] else "",
                    f"<<COMPLETED_WORKORDER{i}_RECOMMENDATION2>>": data['recommendations'].split('\n')[1] if data['recommendations'] and len(data['recommendations'].split('\n')) > 1 else "",
                    f"<<COMPLETED_WORKORDER{i}_RECOMMENDATION3>>": data['recommendations'].split('\n')[2] if data['recommendations'] and len(data['recommendations'].split('\n')) > 2 else "",
                })
                
                # Add photo placeholders if available
                if data['photos']:
                    for j, photo in enumerate(data['photos'][:2], 1):  # Only first 2 photos
                        replacements.update({
                            f"<<COMPLETED_WORKORDER{i}_PHOTO{j}_DESCRIPTION>>": photo['description'],
                            f"<<COMPLETED_WORKORDER{i}_PHOTO{j}_PATH>>": photo['path'] if photo['path'] else "",
                        })
            
            def replace_text_in_paragraph(paragraph, replacements):
                """Replace text while preserving formatting"""
                if not paragraph.runs:
                    return
                
                # Store initial formatting from first run
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                italic = first_run.font.italic
                
                # Get paragraph text and perform replacements
                text = paragraph.text
                for key, value in replacements.items():
                    if key in text:
                        text = text.replace(key, str(value))
                
                # Clear paragraph and add new run with preserved formatting
                paragraph.clear()
                run = paragraph.add_run(text)
                run.font.name = font_name
                run.font.size = font_size
                run.font.bold = bold
                run.font.italic = italic
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if any(key in paragraph.text for key in replacements):
                    replace_text_in_paragraph(paragraph, replacements)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if any(key in paragraph.text for key in replacements):
                                replace_text_in_paragraph(paragraph, replacements)
            
            # Handle photos
            for i, section in enumerate(self.completed_sections, 1):
                data = section.get_data()
                if data['photos']:
                    for j, photo in enumerate(data['photos'][:2], 1):  # Only first 2 photos
                        if photo['path']:
                            # Find placeholder paragraph
                            for paragraph in doc.paragraphs:
                                if f"<<COMPLETED_WORKORDER{i}_PHOTO{j}_PATH>>" in paragraph.text:
                                    # Clear the paragraph
                                    paragraph.clear()
                                    # Add the image
                                    run = paragraph.add_run()
                                    run.add_picture(photo['path'], width=Inches(6))
                                    break
            
            # Save changes
            doc.save(file_path)
            messagebox.showinfo("Success", "Document generated successfully!")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))

def main():
    app = FormFillerApp()
    app.run()

if __name__ == "__main__":
    main()
